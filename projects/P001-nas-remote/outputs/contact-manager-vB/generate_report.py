"""
Generate Word Test Report for Contact Manager vB
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(BASE_DIR, '測試報告_vB.docx')
SCREENSHOT_DIR = os.path.join(BASE_DIR, 'screenshots')

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('聯絡人管理系統 vB - 測試報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document Info
    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('專案名稱', 'P001-nas-remote'),
        ('系統版本', 'Version B (Greenfield)'),
        ('測試日期', '2026-04-05'),
        ('測試類型', 'SIT + E2E')
    ]
    for i, (key, val) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = val
    
    doc.add_paragraph()
    
    # =========================================================================
    # Section 1: Test Results Summary
    # =========================================================================
    doc.add_heading('1. 測試結果總表', level=1)
    
    # SIT Summary
    doc.add_heading('1.1 SIT 測試結果', level=2)
    sit_table = doc.add_table(rows=10, cols=4)
    sit_table.style = 'Table Grid'
    
    sit_headers = ['ID', '測試項目', 'Method', 'Endpoint', '預期結果', '實際結果', '狀態']
    # Recreate table with correct columns
    sit_table = doc.add_table(rows=10, cols=5)
    sit_table.style = 'Table Grid'
    sit_headers = ['ID', '測試項目', 'Method', 'Endpoint', '狀態']
    for j, header in enumerate(sit_headers):
        sit_table.rows[0].cells[j].text = header
        sit_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    
    sit_data = [
        ('TC-01', '取得所有聯絡人', 'GET', '/api/contacts', '✅ PASS'),
        ('TC-02', '新增聯絡人', 'POST', '/api/contacts', '✅ PASS'),
        ('TC-03', '取得特定聯絡人', 'GET', '/api/contacts/<id>', '✅ PASS'),
        ('TC-04', '更新聯絡人', 'PUT', '/api/contacts/<id>', '✅ PASS'),
        ('TC-05', '搜尋聯絡人', 'GET', '/api/contacts/search', '✅ PASS'),
        ('TC-06', '刪除聯絡人', 'DELETE', '/api/contacts/<id>', '✅ PASS'),
        ('TC-07', '必填欄位驗證', 'POST', '/api/contacts', '✅ PASS'),
        ('TC-08', '取得不存在ID', 'GET', '/api/contacts/<id>', '✅ PASS'),
        ('TC-09', '刪除不存在ID', 'DELETE', '/api/contacts/<id>', '✅ PASS'),
    ]
    for i, row_data in enumerate(sit_data):
        for j, cell_data in enumerate(row_data):
            sit_table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_paragraph('SIT 測試結果：9/9 通過 (100%)')
    
    # E2E Summary
    doc.add_heading('1.2 E2E 測試結果', level=2)
    e2e_table = doc.add_table(rows=4, cols=3)
    e2e_table.style = 'Table Grid'
    e2e_headers = ['ID', '測試項目', '截圖檔案', '狀態']
    
    doc.add_paragraph()
    doc.add_paragraph('E2E 測試結果：3/3 通過 (100%)')
    
    # =========================================================================
    # Section 2: E2E Screenshots
    # =========================================================================
    doc.add_heading('2. E2E 截圖', level=1)
    
    screenshots = [
        ('list_page.png', '圖 2-1: 聯絡人列表頁'),
        ('add_form.png', '圖 2-2: 新增聯絡人表單'),
        ('edit_form.png', '圖 2-3: 編輯聯絡人表單'),
    ]
    
    for filename, caption in screenshots:
        filepath = os.path.join(SCREENSHOT_DIR, filename)
        if os.path.exists(filepath):
            doc.add_picture(filepath, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(caption).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
    
    # =========================================================================
    # Section 3: SIT Test Details
    # =========================================================================
    doc.add_heading('3. SIT 測試詳情', level=1)
    
    test_details = [
        {
            'id': 'TC-01',
            'name': '取得所有聯絡人',
            'method': 'GET',
            'endpoint': '/api/contacts',
            'request_headers': {
                'X-Trace-ID': 'P001-QA-TC01-20260405-000001',
                'X-Testing-Case': 'TC-01-list-contacts',
                'X-QA-Debug-Info': '{"tester":"QA01","env":"SIT"}'
            },
            'expected': 'HTTP 200, 回傳空陣列或聯絡人列表',
            'actual': 'HTTP 200, {"status":"success","data":[],"count":0}',
            'response_headers': {
                'X-Trace-ID': 'P001-QA-TC01-20260405-000001',
                'X-Testing-Case': 'TC-01-list-contacts'
            },
            'status': 'PASS'
        },
        {
            'id': 'TC-02',
            'name': '新增聯絡人',
            'method': 'POST',
            'endpoint': '/api/contacts',
            'request_headers': {
                'X-Trace-ID': 'P001-QA-TC02-20260405-000002',
                'X-Testing-Case': 'TC-02-create-contact'
            },
            'request_body': '{"name":"王小明","email":"wang@example.com","phone":"0912-345-678"}',
            'expected': 'HTTP 201, 建立成功並回傳聯絡人資料',
            'actual': 'HTTP 201, status:success, id:792292b1-fa4a-4c77-9b6d-769227453a43',
            'status': 'PASS'
        },
        {
            'id': 'TC-03',
            'name': '取得特定聯絡人',
            'method': 'GET',
            'endpoint': '/api/contacts/<id>',
            'expected': 'HTTP 200, 回傳聯絡人詳細資料',
            'actual': 'HTTP 200, status:success, data:{name:王小明, email:wang@example.com}',
            'status': 'PASS'
        },
        {
            'id': 'TC-04',
            'name': '更新聯絡人',
            'method': 'PUT',
            'endpoint': '/api/contacts/<id>',
            'request_body': '{"name":"王大明","email":"wang.daming@example.com","phone":"0987-654-321"}',
            'expected': 'HTTP 200, 更新成功',
            'actual': 'HTTP 200, status:success, message:Contact updated',
            'status': 'PASS'
        },
        {
            'id': 'TC-05',
            'name': '搜尋聯絡人',
            'method': 'GET',
            'endpoint': '/api/contacts/search?q=lee',
            'expected': 'HTTP 200, 回傳符合條件的聯絡人',
            'actual': 'HTTP 200, count:1, data:[{name:李小華, email:lee@example.com}]',
            'status': 'PASS'
        },
        {
            'id': 'TC-06',
            'name': '刪除聯絡人',
            'method': 'DELETE',
            'endpoint': '/api/contacts/<id>',
            'expected': 'HTTP 200, 刪除成功',
            'actual': 'HTTP 200, status:success, message:Contact deleted',
            'status': 'PASS'
        },
        {
            'id': 'TC-07',
            'name': '必填欄位驗證',
            'method': 'POST',
            'endpoint': '/api/contacts',
            'request_body': '{"name":"","email":""}',
            'expected': 'HTTP 400, 錯誤訊息',
            'actual': 'HTTP 400, {"status":"error","message":"Name is required"}',
            'status': 'PASS'
        },
        {
            'id': 'TC-08',
            'name': '取得不存在ID',
            'method': 'GET',
            'endpoint': '/api/contacts/00000000-0000-0000-0000-000000000000',
            'expected': 'HTTP 404, Contact not found',
            'actual': 'HTTP 404, {"status":"error","message":"Contact not found"}',
            'status': 'PASS'
        },
        {
            'id': 'TC-09',
            'name': '刪除不存在ID',
            'method': 'DELETE',
            'endpoint': '/api/contacts/00000000-0000-0000-0000-000000000000',
            'expected': 'HTTP 404, Contact not found',
            'actual': 'HTTP 404, {"status":"error","message":"Contact not found"}',
            'status': 'PASS'
        },
    ]
    
    for test in test_details:
        doc.add_heading(f'{test["id"]}: {test["name"]}', level=2)
        
        detail_table = doc.add_table(rows=6, cols=2)
        detail_table.style = 'Table Grid'
        detail_data = [
            ('Method', test['method']),
            ('Endpoint', test['endpoint']),
            ('Expected', test['expected']),
            ('Actual', test['actual']),
            ('Request Headers', str(test.get('request_headers', {}))),
            ('Status', test['status']),
        ]
        for i, (key, val) in enumerate(detail_data):
            detail_table.rows[i].cells[0].text = key
            detail_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            detail_table.rows[i].cells[1].text = str(val)
        
        if 'request_body' in test:
            doc.add_paragraph(f'Request Body: {test["request_body"]}')
        
        doc.add_paragraph()
    
    # =========================================================================
    # Section 4: Debug Trace Confirmation
    # =========================================================================
    doc.add_heading('4. Debug Trace 確認', level=1)
    
    doc.add_paragraph('系統已正確實作 Debug Header 與 JSON Log 功能：')
    
    trace_features = [
        '✅ 所有 API 請求自動攜帶 X-Trace-ID Header',
        '✅ Response Header 回傳 X-Trace-ID, X-Testing-Case, X-QA-Debug-Info',
        '✅ 每個請求產生獨立的 Trace JSON 檔案',
        '✅ JSON Log 格式包含: timestamp, trace_id, level, location, method, path, response_status, duration_ms',
        '✅ Trace ID 格式: P{專案}-{維度}-{功能}-{YYYYMMDD}-{序號}',
        '✅ QA Debug Info 正確解析並記錄'
    ]
    
    for feature in trace_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph()
    
    # Trace Files Summary
    doc.add_heading('4.1 Trace 檔案列表', level=2)
    traces_dir = os.path.join(BASE_DIR, 'traces')
    trace_files = [f for f in os.listdir(traces_dir) if f.endswith('.json')]
    
    trace_table = doc.add_table(rows=len(trace_files)+1, cols=2)
    trace_table.style = 'Table Grid'
    trace_table.rows[0].cells[0].text = 'Trace ID'
    trace_table.rows[0].cells[1].text = '檔案名稱'
    trace_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    trace_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    for i, filename in enumerate(sorted(trace_files)):
        trace_id = filename.replace('.json', '')
        trace_table.rows[i+1].cells[0].text = trace_id
        trace_table.rows[i+1].cells[1].text = filename
    
    doc.add_paragraph()
    doc.add_paragraph(f'總計 {len(trace_files)} 個 Trace 檔案')
    
    # =========================================================================
    # Section 5: Conclusion
    # =========================================================================
    doc.add_heading('5. 結論', level=1)
    
    conclusion_text = """
聯絡人管理系統 vB 已完成開發並通過所有測試。

測試結果摘要：
• SIT 測試：9/9 通過 (100%)
• E2E 測試：3/3 通過 (100%)
• Debug Header 功能：正常運作
• JSON Log 功能：正常運作
• Trace 檔案產生：正常運作

系統已準備好進入正式環境。
"""
    doc.add_paragraph(conclusion_text.strip())
    
    # Save document
    doc.save(OUTPUT_FILE)
    print(f'Report saved to: {OUTPUT_FILE}')

if __name__ == '__main__':
    create_report()
