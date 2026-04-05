from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc_path = "/home/devpro/Claude-Workspace/projects/P001-nas-remote/outputs/contact-manager/P001_測試報告_20260405.docx"
doc = Document()

# Title
title = doc.add_heading('📇 P001 聯絡人管理系統 - 測試報告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Project Info
doc.add_paragraph('專案位置：/home/devpro/Claude-Workspace/projects/P001-nas-remote/outputs/contact-manager/')
doc.add_paragraph('測試時間：2026-04-05 18:19 GMT+8')
doc.add_paragraph('Flask 伺服器：http://localhost:5000')
doc.add_paragraph('執行者：QA Engineer (Subagent)')
doc.add_paragraph()

# Summary Table
doc.add_heading('📊 測試結果總表', level=1)
summary_data = [
    ('維度', '項目', '結果', '備註'),
    ('SIT', 'P001-SIT-001: GET /api/contacts', '✅ Pass', '正常回傳 2 筆聯絡人'),
    ('SIT', 'P001-SIT-002: GET /api/contacts/<id>', '✅ Pass', '正常回傳指定聯絡人'),
    ('SIT', 'P001-SIT-003: POST /api/contacts', '✅ Pass', '建立成功，回傳 201'),
    ('SIT', 'P001-SIT-004: PUT /api/contacts/<id>', '✅ Pass', '更新成功，欄位正確'),
    ('SIT', 'P001-SIT-005: DELETE /api/contacts/<id>', '✅ Pass', '刪除成功，回傳 200'),
    ('SIT', 'P001-SIT-006: GET /api/contacts/search', '✅ Pass', '搜尋「小」回傳 1 筆'),
    ('SIT', 'P001-SIT-007: 必填欄位驗證', '✅ Pass', '未提供 name 回傳 400'),
    ('SIT', 'P001-SIT-008: 取得不存在的 ID', '✅ Pass', '回傳 404'),
    ('SIT', 'P001-SIT-009: 刪除不存在的 ID', '✅ Pass', '回傳 404'),
    ('E2E', 'P001-E2E-001: 列表頁載入', '✅ Pass', '截圖已存'),
    ('E2E', 'P001-E2E-002: 新增表單頁載入', '✅ Pass', '截圖已存'),
    ('E2E', 'P001-E2E-003: 編輯表單頁載入', '✅ Pass', '截圖已存'),
    ('UT', 'P001-UT-001: app.py 函式邏輯', '✅ Pass', 'CRUD 邏輯正確'),
    ('UT', 'P001-UT-002: UUID 生成', '✅ Pass', '使用 uuid.uuid4()'),
    ('UT', 'P001-UT-003: JSON 讀寫', '✅ Pass', 'utf-8 編碼正確'),
    ('UT', 'P001-UT-004: 搜尋比對', '✅ Pass', '大小寫不敏感比對正確'),
]

table = doc.add_table(rows=len(summary_data), cols=4)
table.style = 'Table Grid'
for i, row_data in enumerate(summary_data):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        if i == 0:  # Header row
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('總計：16/16 項全部通過')
run.bold = True

doc.add_paragraph()

# SIT Details
doc.add_heading('SIT（系統整合測試）- 詳情', level=1)
sit_details = [
    'P001-SIT-001: GET /api/contacts\n  預期：回傳所有聯絡人 | 實際：✅ 回傳 JSON 陣列，含 2 筆資料（HTTP 200）',
    'P001-SIT-002: GET /api/contacts/<id>\n  預期：回傳指定聯絡人 | 實際：✅ 回傳完整聯絡人 JSON（HTTP 200）',
    'P001-SIT-003: POST /api/contacts\n  預期：建立新聯絡人，回傳 201 | 實際：✅ 建立成功，回傳新 UUID（HTTP 201）',
    'P001-SIT-004: PUT /api/contacts/<id>\n  預期：更新聯絡人 | 實際：✅ 更新成功（HTTP 200）',
    'P001-SIT-005: DELETE /api/contacts/<id>\n  預期：刪除聯絡人 | 實際：✅ 回傳 {"message":"Contact deleted"}（HTTP 200）',
    'P001-SIT-006: GET /api/contacts/search?q=小\n  預期：搜尋回傳符合的聯絡人 | 實際：✅ 回傳 1 筆（小明）（HTTP 200）',
    'P001-SIT-007: 必填欄位驗證\n  預期：回傳 400 錯誤 | 實際：✅ {"error":"Name is required"}（HTTP 400）',
    'P001-SIT-008: GET /api/contacts/fake-id\n  預期：回傳 404 | 實際：✅ {"error":"Contact not found"}（HTTP 404）',
    'P001-SIT-009: DELETE /api/contacts/fake-id\n  預期：回傳 404 | 實際：✅ {"error":"Contact not found"}（HTTP 404）',
]
for detail in sit_details:
    doc.add_paragraph(detail)

doc.add_paragraph()

# E2E Details
doc.add_heading('E2E（端對端測試）- 詳情', level=1)
e2e_details = [
    'P001-E2E-001: 列表頁載入（http://localhost:5000/）\n  結果：✅ Pass | 截圖：P001-E2E-001-list-20260405_181902.png',
    'P001-E2E-002: 新增表單頁載入（http://localhost:5000/contacts/new）\n  結果：✅ Pass | 截圖：P001-E2E-002-new-20260405_181902.png',
    'P001-E2E-003: 編輯表單頁載入\n  結果：✅ Pass | 截圖：P001-E2E-003-edit-20260405_181902.png',
]
for detail in e2e_details:
    doc.add_paragraph(detail)

doc.add_paragraph()

# UT Details
doc.add_heading('UT（單元測試）- 程式碼審查', level=1)
ut_details = [
    'P001-UT-001: app.py 函式邏輯\n  ✅ CRUD 四個 API 路由邏輯正確，錯誤處理（404/400）正確',
    'P001-UT-002: UUID 生成\n  ✅ 使用 uuid.uuid4() 產生全球唯一 ID',
    'P001-UT-003: JSON 讀寫\n  ✅ utf-8 編碼正確，ensure_ascii=False 支援中文',
    'P001-UT-004: 搜尋比對\n  ✅ 大小寫不敏感比對正確，搜尋欄位：name、phone、email、company、notes',
]
for detail in ut_details:
    doc.add_paragraph(detail)

doc.add_paragraph()

# Issues
doc.add_heading('🔍 發現的問題', level=1)
p = doc.add_paragraph()
run = p.add_run('無重大問題。系統運作正常，所有測試項目皆通過。')
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph('輕微觀察（非 Bug）：')
doc.add_paragraph('1. data/contacts.json 目前只有 2 筆測試資料，可考慮增加邊界條件測試')
doc.add_paragraph('2. DELETE 成功後不回傳 204，而是回傳 200 + message（屬於風格問題，不影響功能）')

doc.add_paragraph()

# Footer
footer = doc.add_paragraph('測試完成時間：2026-04-05 18:19 GMT+8')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('測試結果：✅ 16/16 全部通過')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(doc_path)
print(f'✅ 報告已產生：{doc_path}')