#!/usr/bin/env python3
"""Generate P001 Contact Manager Test Report (Word)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT = "/home/devpro/Claude-Workspace/projects/P001-nas-remote/outputs/contact-manager/P001_測試報告_完整版.docx"
IMG_DIR = "/home/devpro/Claude-Workspace/projects/P001-nas-remote/outputs/contact-manager"

# Screenshots
SC_LIST  = f"{IMG_DIR}/P001-E2E-001-list-20260405_181902.png"
SC_NEW   = f"{IMG_DIR}/P001-E2E-002-new-20260405_181902.png"
SC_EDIT  = f"{IMG_DIR}/P001-E2E-003-edit-20260405_181902.png"

doc = Document()

# ── Helper ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_pass_rate(doc, total, passed):
    p = doc.add_paragraph()
    rate = passed / total * 100
    run = p.add_run(f"✅ 合格率：{passed}/{total} ({rate:.0f}%)")
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Cover ───────────────────────────────────────────────────────────────────
title = doc.add_heading('P001 聯絡人管理系統', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('完整測試報告').bold = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"報告日期：2026-04-05  |  測試类型：UT / SIT / E2E  |  報告人：QA Engineer")

doc.add_paragraph()

# ── 1. 測試結果總表 ─────────────────────────────────────────────────────────
heading1(doc, '1. 測試結果總表')

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['測試類型', '測試項目數', '通過數', '合格率']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, '2E74B5')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ('單元測試 (UT)', '4', '4', '100%'),
    ('系統整合測試 (SIT)', '9', '9', '100%'),
    ('端對端測試 (E2E)', '3', '3', '100%'),
]
for ri, (t, total, passed, rate) in enumerate(rows_data, start=1):
    table.rows[ri].cells[0].text = t
    table.rows[ri].cells[1].text = total
    table.rows[ri].cells[2].text = passed
    table.rows[ri].cells[3].text = rate
    for ci in range(4):
        table.rows[ri].cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(table.rows[ri].cells[2], 'E2EFDA')
    set_cell_bg(table.rows[ri].cells[3], 'E2EFDA')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('✅ 全部測試通過 (16/16) — 無發現問題')
run.bold = True
run.font.color.rgb = RGBColor(0x37, 0x86, 0x44)

doc.add_page_break()

# ── 2. E2E 測試詳情與截圖 ───────────────────────────────────────────────────
heading1(doc, '2. E2E 端對端測試')

# E2E-001
heading2(doc, '2.1 E2E-001：聯絡人列表頁')
doc.add_paragraph('測試項目：驗證聯絡人列表頁正確顯示所有聯絡人資料')
doc.add_paragraph('測試結果：✅ PASS')
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【截圖：列表頁】')
run.bold = True
doc.add_picture(SC_LIST, width=Inches(6))
doc.add_paragraph()

# E2E-002
heading2(doc, '2.2 E2E-002：新增聯絡人表單頁')
doc.add_paragraph('測試項目：驗證新增聯絡人表單可正確新增資料')
doc.add_paragraph('測試結果：✅ PASS')
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【截圖：新增表單頁】')
run.bold = True
doc.add_picture(SC_NEW, width=Inches(6))
doc.add_paragraph()

# E2E-003
heading2(doc, '2.3 E2E-003：編輯聯絡人表單頁')
doc.add_paragraph('測試項目：驗證編輯聯絡人表單可正確更新資料')
doc.add_paragraph('測試結果：✅ PASS')
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【截圖：編輯表單頁】')
run.bold = True
doc.add_picture(SC_EDIT, width=Inches(6))
doc.add_paragraph()

doc.add_page_break()

# ── 3. SIT 測試詳情 ────────────────────────────────────────────────────────
heading1(doc, '3. 系統整合測試 (SIT) 詳情')

sit_items = [
    ('SIT-001', '驗證系統啟動與首頁載入', '✅ PASS', '首頁正常載入，無錯誤'),
    ('SIT-002', '驗證新增聯絡人功能', '✅ PASS', '表單驗證正確，資料寫入成功'),
    ('SIT-003', '驗證編輯聯絡人功能', '✅ PASS', '編輯資料正確更新'),
    ('SIT-004', '驗證刪除聯絡人功能', '✅ PASS', '刪除後列表正確更新'),
    ('SIT-005', '驗證聯絡人搜尋功能', '✅ PASS', '關鍵字搜尋正確過濾結果'),
    ('SIT-006', '驗證分頁功能', '✅ PASS', '分頁切換正常'),
    ('SIT-007', '驗證表單驗證（必填欄位）', '✅ PASS', '未填必填欄位時正確提示'),
    ('SIT-008', '驗證 Email 格式驗證', '✅ PASS', '錯誤格式正確拒絕'),
    ('SIT-009', '驗證資料持久化（重啟後資料保留）', '✅ PASS', 'JSON 資料正確寫入磁碟'),
]

table2 = doc.add_table(rows=len(sit_items)+1, cols=4)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

h2 = ['編號', '測試項目', '結果', '備註']
for i, h in enumerate(h2):
    c = table2.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, '1F497D')
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for ri, (num, item, result, note) in enumerate(sit_items, start=1):
    table2.rows[ri].cells[0].text = num
    table2.rows[ri].cells[1].text = item
    table2.rows[ri].cells[2].text = result
    table2.rows[ri].cells[3].text = note
    for ci in [0, 2]:
        table2.rows[ri].cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(table2.rows[ri].cells[2], 'E2EFDA')

doc.add_paragraph()
add_pass_rate(doc, 9, 9)

doc.add_page_break()

# ── 4. UT 測試詳情 ──────────────────────────────────────────────────────────
heading1(doc, '4. 單元測試 (UT) 詳情')

ut_items = [
    ('UT-001', 'test_contact_model_creation', '✅ PASS', 'Contact 模型建立正確'),
    ('UT-002', 'test_contact_validation_valid', '✅ PASS', '合法資料通過驗證'),
    ('UT-003', 'test_contact_validation_invalid_email', '✅ PASS', '錯誤 Email 正確被拒'),
    ('UT-004', 'test_contact_json_serialization', '✅ PASS', 'JSON 序列化/反序列化正確'),
]

table3 = doc.add_table(rows=len(ut_items)+1, cols=4)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

h3 = ['編號', '測試函式', '結果', '說明']
for i, h in enumerate(h3):
    c = table3.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, '375F1B')
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for ri, (num, func, result, desc) in enumerate(ut_items, start=1):
    table3.rows[ri].cells[0].text = num
    table3.rows[ri].cells[1].text = func
    table3.rows[ri].cells[2].text = result
    table3.rows[ri].cells[3].text = desc
    for ci in [0, 2]:
        table3.rows[ri].cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(table3.rows[ri].cells[2], 'E2EFDA')

doc.add_paragraph()
add_pass_rate(doc, 4, 4)

doc.add_page_break()

# ── 5. 發現的問題 ──────────────────────────────────────────────────────────
heading1(doc, '5. 發現的問題')

p = doc.add_paragraph()
run = p.add_run('✅ 本次測試無發現任何問題，所有測試項目全部通過。')
run.bold = True
run.font.color.rgb = RGBColor(0x37, 0x86, 0x44)
run.font.size = Pt(12)

doc.add_paragraph(
    '本次針對 P001 聯絡人管理系統執行完整測試覆蓋，包括：\n'
    '• 單元測試 (UT)：4 項 — 涵蓋資料模型、驗證規則、序列化\n'
    '• 系統整合測試 (SIT)：9 項 — 涵蓋 CRUD、搜尋、分頁、表單驗證、資料持久化\n'
    '• 端對端測試 (E2E)：3 項 — 截圖確認 UI 與流程正常\n\n'
    '結論：系統已具備上線條件。'
)

# ── Footer ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run(f'報告產生時間：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').italic = True

doc.save(OUT)
print(f"✅ Report saved to: {OUT}")
