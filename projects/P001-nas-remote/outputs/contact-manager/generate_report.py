from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

output_dir = "/home/devpro/Claude-Workspace/projects/P001-nas-remote/outputs/contact-manager"
doc_path = f"{output_dir}/測試報告_P001_聯絡人管理系統.docx"

doc = Document()

# Title
title = doc.add_heading('📇 聯絡人管理系統 - 測試報告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Project info
doc.add_paragraph('專案：P001 - 聯絡人管理系統（網頁版）')
doc.add_paragraph('測試日期：2026-04-05 17:50 GMT+8')
doc.add_paragraph('測試人員：龍哥（QA Engineer）')
doc.add_paragraph()

# Summary
doc.add_heading('📊 測試摘要', level=1)
summary_table = doc.add_table(rows=6, cols=2)
summary_table.style = 'Table Grid'
summary_data = [
    ('功能', '結果'),
    ('後端 API - GET /api/contacts', '✅ 通過'),
    ('後端 API - POST /api/contacts', '✅ 通過'),
    ('後端 API - DELETE /api/contacts', '✅ 通過'),
    ('前端頁面 - 列表頁', '✅ 通過'),
    ('前端頁面 - 新增表單', '✅ 通過'),
]
for i, (col1, col2) in enumerate(summary_data):
    row = summary_table.rows[i]
    row.cells[0].text = col1
    row.cells[1].text = col2

doc.add_paragraph()

# API Test Details
doc.add_heading('🔧 API 功能測試', level=1)

api_tests = [
    ('GET /api/contacts', '✅', '初始空陣列 [] 正常返回'),
    ('POST /api/contacts', '✅', '成功新增資料，UUID 自動生成'),
    ('GET /api/contacts（再次）', '✅', '正確顯示已新增的資料'),
    ('DELETE /api/contacts', '✅', '刪除成功，返回訊息'),
    ('GET /api/contacts（刪除後）', '✅', '確認資料已移除'),
]

for endpoint, result, desc in api_tests:
    p = doc.add_paragraph()
    p.add_run(f'{result} {endpoint}: ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# Frontend Test
doc.add_heading('🖥️ 前端頁面測試', level=1)

doc.add_paragraph('測試方式：使用 Playwright Headless Chromium 截圖')

doc.add_paragraph()
p = doc.add_paragraph('截圖 1：聯絡人列表頁')
p.runs[0].bold = True
try:
    doc.add_picture(f'{output_dir}/test_screenshot_01_list.png', width=Inches(6))
    doc.add_paragraph('✅ 列表頁正常載入，包含搜尋列和新增按鈕')
except:
    doc.add_paragraph('⚠️ 截圖無法嵌入')

doc.add_paragraph()
p = doc.add_paragraph('截圖 2：新增聯絡人表單')
p.runs[0].bold = True
try:
    doc.add_picture(f'{output_dir}/test_screenshot_02_add_form.png', width=Inches(6))
    doc.add_paragraph('✅ 新增表單正常載入，包含所有欄位')
except:
    doc.add_paragraph('⚠️ 截圖無法嵌入')

doc.add_paragraph()

# Conclusion
doc.add_heading('📝 結論', level=1)
doc.add_paragraph('✅ 所有測試項目均已通過')
doc.add_paragraph('✅ 系統具備基本功能，可進入部署階段')
doc.add_paragraph('✅ 前端頁面呈現正常，無明顯錯誤')

doc.add_paragraph()

# Footer
footer = doc.add_paragraph('報告產生時間：2026-04-05 17:50:00')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(doc_path)
print(f"✅ 報告已產生：{doc_path}")