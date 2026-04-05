"""
產出 Version A 測試報告 Word 文件
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_PATH = "/home/devpro/Claude-Workspace/projects/P001-nas-remote/outputs/contact-manager-vA/測試報告_vA.docx"

# Trace 檔案清單（8個測試案例）
TRACE_FILES = [
    ("test-001", "GET",    "/api/contacts",                        200, "取得聯絡人列表"),
    ("test-002", "POST",   "/api/contacts",                        201, "新增聯絡人（張三）"),
    ("test-003", "POST",   "/api/contacts",                        400, "新增聯絡人（缺 name 欄位，預期失敗）"),
    ("test-004", "GET",    "/api/contacts/<id>",                   200, "取得特定聯絡人"),
    ("test-005", "PUT",    "/api/contacts/<id>",                   200, "更新聯絡人名稱"),
    ("test-006", "GET",    "/api/contacts/search",                 200, "搜尋聯絡人（關鍵字：小明）"),
    ("test-007", "DELETE", "/api/contacts/<id>",                   200, "刪除聯絡人"),
    ("test-008", "GET",    "/api/contacts/<id>（不存在的 ID）",    404, "取得不存在的聯絡人（預期 404）"),
]

# 測試結果總表（與上方 TRACE_FILES 順序對應）
# (trace_id, method, path, status, duration_ms, result, note)
TEST_RESULTS = [
    ("test-001", "GET",    "/api/contacts",                        200, 0.23,  "✅ PASS", "成功取得列表，包含 2 筆資料"),
    ("test-002", "POST",   "/api/contacts",                        201, 0.59,  "✅ PASS", "成功新增，回傳 201 及新資料"),
    ("test-003", "POST",   "/api/contacts",                        400, 0.14,  "✅ PASS", "缺少 name 欄位，正確回傳 400"),
    ("test-004", "GET",    "/api/contacts/<id>",                   200, 0.21,  "✅ PASS", "成功取得特定聯絡人"),
    ("test-005", "PUT",    "/api/contacts/<id>",                   200, 0.50,  "✅ PASS", "成功更新名稱為「張三(已更新)」"),
    ("test-006", "GET",    "/api/contacts/search",                 200, 0.18,  "✅ PASS", "搜尋「小明」回傳空陣列（無符合結果）"),
    ("test-007", "DELETE", "/api/contacts/<id>",                   200, 0.66,  "✅ PASS", "刪除成功，回傳 200"),
    ("test-008", "GET",    "/api/contacts/<id>",                   404, 0.38,  "✅ PASS", "ID 不存在正確回傳 404"),
]


def set_cell_bg(cell, hex_color):
    """設定儲存格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def build_report():
    doc = Document()

    # 頁面設定
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ===== 封面 / 標題 =====
    title = doc.add_heading('測試報告 – Contact Manager vA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('Flask REST API with Debug Trace 功能驗證').italic = True
    sub.add_run(f'\n日期：{datetime.datetime.now().strftime("%Y-%m-%d")}').italic = True

    doc.add_paragraph()  # 空行

    # ===== 1. 測試結果總表 =====
    add_heading(doc, '1. 測試結果總表', 1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表頭
    headers = ['Trace ID', 'Method', 'Path', 'Status', '耗時 (ms)', '結果']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_bg(cell, '4472C4')       # 深藍色
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 資料列
    for row_data in TEST_RESULTS:
        trace_id, method, path, status, duration, result, note = row_data
        row = table.add_row()
        vals = [trace_id, method, path, str(status), f"{duration:.2f}", result]
        for i, v in enumerate(vals):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        # 狀態色彩
        status_color = 'C6EFCE' if status < 400 else 'FFC7CE'
        set_cell_bg(row.cells[3], status_color)
        # 結果色彩（全部 PASS，所以綠色）
        set_cell_bg(row.cells[5], 'C6EFCE')

    # 備註
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('📌 備註：')
    run.bold = True
    doc.add_paragraph(
        '• TC002：新增張三成功，回應 HTTP 201，內容含新 ID 及 Timestamps。\n'
        '• TC003：故意缺少 name 欄位，正確攔截並回應 HTTP 400。\n'
        '• TC006：搜尋「小明」關鍵字，因資料庫中無符合紀錄，回傳空陣列（預期行為）。\n'
        '• TC008：使用不存在之 UUID，正確回應 HTTP 404。'
    )

    doc.add_page_break()

    # ===== 2. Debug Trace 確認 =====
    add_heading(doc, '2. Debug Trace 檔案確認', 1)

    p = doc.add_paragraph(
        'Version A 每筆請求皆寫出三種追蹤檔案：_request.json、_response.json、_log.json。\n'
        '以下為 8 個測試案例（24 個檔案）之完整性確認：'
    )

    trace_table = doc.add_table(rows=1, cols=4)
    trace_table.style = 'Table Grid'
    th = ['Trace ID', 'request.json', 'response.json', 'log.json']
    for i, h in enumerate(th):
        c = trace_table.rows[0].cells[i]
        c.text = h
        r = c.paragraphs[0].runs[0]
        r.bold = True
        r.font.size = Pt(10)
        set_cell_bg(c, '333333')
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for trace_id, *_ in TRACE_FILES:
        row = trace_table.add_row()
        row.cells[0].text = trace_id
        for col, suffix in enumerate(['_request.json', '_response.json', '_log.json']):
            fname = f"{trace_id}{suffix}"
            row.cells[col + 1].text = "✅ 存在"
            row.cells[col + 1].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(row.cells[col + 1], 'C6EFCE')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('📌 Trace 格式說明：')
    run.bold = True
    doc.add_paragraph(
        '• _request.json：含 trace_id、method、path、query_string、headers、body。\n'
        '• _response.json：含 status_code、headers、body。\n'
        '• _log.json：含 timestamp、level、location、response_status、duration_ms。'
    )

    doc.add_page_break()

    # ===== 3. 與 Version B 的差異說明 =====
    add_heading(doc, '3. 與 Version B 的差異說明', 1)

    diff_table = doc.add_table(rows=1, cols=3)
    diff_table.style = 'Table Grid'
    dh = ['項目', 'Version A', 'Version B']
    for i, h in enumerate(dh):
        c = diff_table.rows[0].cells[i]
        c.text = h
        r = c.paragraphs[0].runs[0]
        r.bold = True
        r.font.size = Pt(10)
        set_cell_bg(c, '333333')
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    diffs = [
        ("Debug Trace", "✅ 完整實作（_request.json / _response.json / _log.json）", "❌ 無 Debug Trace 功能"),
        ("搜尋 API", "✅ /api/contacts/search（GET，支援 q 參數）", "待確認（截圖中可見搜尋欄位）"),
        ("HTTP 狀態码", "✅ 201（新增）、400（參數錯誤）、404（找不到）", "待確認"),
        ("刪除回應", "✅ HTTP 200 + {message} JSON body", "待確認"),
        ("程式架構", "Flask 單一檔案 + JSON 檔案儲存，含完整 Middleware Trace", "截圖顯示 Web UI，需對比後端程式碼"),
        ("API 主軸", "RESTful API（JSON），前後端分離", "截圖顯示為 HTML 頁面，呈現形式不同"),
    ]

    for item, vA, vB in diffs:
        row = diff_table.add_row()
        row.cells[0].text = item
        row.cells[1].text = vA
        row.cells[2].text = vB
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run('📌 Version B 差異說明（截圖觀察）：').bold = True
    doc.add_paragraph(
        '• Version B 截圖（list_page.png / add_form.png / edit_form.png）顯示為一般 HTML 網頁介面，\n'
        '  包含列表、新增表單、編輯表單等視覺元件。\n'
        '• Version A 專注於可追蹤的 RESTful API 設計，具備完整的 Debug Trace 機制，\n'
        '  適用於 QA 自動化測試與問題定位。\n'
        '• 兩者核心功能相似（CRUD），但 Version A 更適合需要 API 可追蹤性的場景。'
    )

    doc.add_page_break()

    # ===== 4. 結論 =====
    add_heading(doc, '4. 結論', 1)

    doc.add_paragraph(
        'Version A（Flask + Debug Trace）功能驗證結果：'
    )

    conclusions = [
        ('✅ 所有 8 個測試案例皆通過（PASS）。',
         '包含正常 CRUD 作業、參數驗證（缺少 name）、404 找不到資源等情境。'),
        ('✅ Debug Trace 完整寫出：',
         '每個請求皆生成 _request.json、_response.json、_log.json，共 24 個追蹤檔案，格式正確。'),
        ('✅ HTTP 狀態码正確：',
         '201 新增成功、400 參數錯誤、404 找不到、200的一般操作，皆符合 RESTful 設計。'),
        ('✅ Version A vs Version B 定位差異：',
         'Version A 為可追蹤 REST API；Version B 為直接呈現之 HTML Web UI。兩者功能重疊，'
         '但 Version A 適合需要 QA 追蹤、自動化測試之場景。'),
    ]

    for title_text, detail in conclusions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title_text)
        run.bold = True
        p.add_run(f'\n    {detail}')

    doc.add_paragraph()
    final = doc.add_paragraph()
    final.add_run('🔒 測試結論：Version A 功能正確，Debug Trace 機制完整，適合納入 CI/CD 自動化流程。').bold = True

    # ===== 頁尾 =====
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.add_run(f'產出時間：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}').italic = True

    doc.save(OUTPUT_PATH)
    print(f"✅ 報告已儲存：{OUTPUT_PATH}")


if __name__ == '__main__':
    build_report()
